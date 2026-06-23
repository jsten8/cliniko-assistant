"""Populates a Word .docx template with field values and converts to PDF."""
from __future__ import annotations
import shutil
from pathlib import Path
from datetime import date
import docx
import config


def _replace_in_paragraph(para, fields: dict[str, str]) -> None:
    for run in para.runs:
        for key, val in fields.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, val)


def _replace_in_table(table, fields: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, fields)


def populate_template(template_path: Path, fields: dict[str, str], output_path: Path) -> Path:
    """
    Copy template, replace all {{key}} placeholders with values, save as .docx.
    Returns the output .docx path.
    """
    shutil.copy2(template_path, output_path)
    doc = docx.Document(output_path)

    for para in doc.paragraphs:
        _replace_in_paragraph(para, fields)

    for table in doc.tables:
        _replace_in_table(table, fields)

    # Also handle headers/footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, fields)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, fields)

    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path: Path) -> Path:
    """
    Convert .docx to PDF. Tries LibreOffice headless first, then MS Word via
    docx2pdf, then falls back to returning the .docx itself so the workflow
    never breaks even without a converter installed.
    """
    import subprocess
    pdf_path = docx_path.with_suffix(".pdf")

    # LibreOffice (free, works headlessly)
    for soffice in ["/Applications/LibreOffice.app/Contents/MacOS/soffice", "soffice", "libreoffice"]:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir",
                 str(docx_path.parent), str(docx_path)],
                capture_output=True, timeout=60,
            )
            if result.returncode == 0 and pdf_path.exists():
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # MS Word via docx2pdf — run with a timeout so it never hangs forever
    try:
        import threading
        from docx2pdf import convert
        result = [None]
        error = [None]

        def _convert():
            try:
                convert(str(docx_path), str(pdf_path))
                result[0] = True
            except Exception as e:
                error[0] = e

        t = threading.Thread(target=_convert, daemon=True)
        t.start()
        t.join(timeout=60)  # Give Word 60s max

        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass

    # Fallback: return the .docx — opens in Pages/Word fine
    return docx_path


def build_output_filename(patient_name: str, workflow_key: str) -> str:
    """Generate a clean output filename."""
    today = date.today().strftime("%d%b%Y")
    safe_name = "".join(c for c in patient_name if c.isalnum() or c in " _").replace(" ", "")
    label_map = {
        "epc_new": "EPC",
        "epc_final": "EPCFinal",
        "dva_new": "DVA",
        "dva_final": "DVAFinal",
        "wc_new": "WC",
        "wc_final_form032": "WCForm032",
        "wc_final_crm": "WCCRMLetter",
    }
    label = label_map.get(workflow_key, workflow_key.upper())
    return f"{label}_{safe_name}_{today}"


def get_template_path(workflow_key: str) -> Path | None:
    """Return the stored template path for a workflow key, or None if not uploaded."""
    path = config.TEMPLATES_DIR / f"{workflow_key}.docx"
    return path if path.exists() else None
